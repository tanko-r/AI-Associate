"""
Rebuild a DOCX file from revision maps.

Takes the original DOCX and one or more revision map JSON files,
produces a clean revised DOCX with all changes applied.

Usage:
    python rebuild_docx.py <original.docx> <output.docx> <revision1.json> [revision2.json ...]
"""
import json
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).parent.parent / "docx-tools"))
from core.inserter import insert_paragraph_after, remove_paragraph
from core.redliner import _build_char_format_map, _make_run_from_rpr


def load_revisions(*json_paths):
    """Load and merge revision maps from multiple JSON files."""
    all_revisions = {}
    for path in json_paths:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        revisions = data.get('revisions', data) if isinstance(data, dict) else data
        if isinstance(revisions, list):
            for rev in revisions:
                rev_id = rev.get('id', rev.get('paragraph_id'))
                all_revisions[float(rev_id)] = rev
        elif isinstance(revisions, dict):
            for k, v in revisions.items():
                all_revisions[float(k)] = v
    return all_revisions


def rebuild_document(original_path, output_path, revisions):
    """Rebuild the document with revisions applied."""
    doc = Document(original_path)

    # Build a map of paragraph index -> paragraph ID (matching extract_paragraphs.py)
    para_id_map = {}  # id -> index in doc.paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            para_id_map[i] = para

    # Collect insertions (keyed by insert_after_id)
    insertions = {}
    for rev_id, rev in sorted(revisions.items()):
        action = rev.get('action', 'revise')
        if action == 'insert_after':
            after_id = rev.get('insert_after_id', int(rev_id))
            if after_id not in insertions:
                insertions[after_id] = []
            insertions[after_id].append(rev)

    # Process revisions on existing paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        rev = revisions.get(float(i))
        if rev and rev.get('action') in ('revise', 'replace'):
            new_text = rev.get('revised_text', '')
            if new_text:
                # Build char map from original runs before any mutations
                char_map = _build_char_format_map(para)

                # Remove all existing <w:r> elements
                p_elem = para._p
                for child in list(p_elem):
                    if child.tag == qn('w:r'):
                        p_elem.remove(child)

                # Rebuild runs mapping revised text to original formatting
                if char_map:
                    # Walk through revised text, assigning formatting from char_map
                    # Characters beyond original length get the last run's formatting
                    current_rpr = None
                    current_run_index = -1
                    segment_start = 0

                    for ci, char in enumerate(new_text):
                        if ci < len(char_map):
                            rpr = char_map[ci].rpr_element
                            run_idx = char_map[ci].run_index
                        else:
                            # Beyond original text -- use last run's formatting
                            rpr = char_map[-1].rpr_element
                            run_idx = char_map[-1].run_index + 1  # force new segment

                        # Detect run boundary change
                        if run_idx != current_run_index:
                            # Emit previous segment
                            if segment_start < ci:
                                new_run = _make_run_from_rpr(new_text[segment_start:ci], current_rpr)
                                p_elem.append(new_run)
                            current_rpr = rpr
                            current_run_index = run_idx
                            segment_start = ci

                    # Emit final segment
                    if segment_start < len(new_text):
                        new_run = _make_run_from_rpr(new_text[segment_start:], current_rpr)
                        p_elem.append(new_run)
                else:
                    # No char map (no original runs) -- just add text with no formatting
                    new_run = _make_run_from_rpr(new_text, None)
                    p_elem.append(new_run)

        elif rev and rev.get('action') == 'delete':
            # Remove the paragraph entirely via XML manipulation
            remove_paragraph(doc, f"p_{i + 1}")

    # Handle insertions — insert at correct positions using XML manipulation
    # Process in reverse order so earlier insertions don't shift later targets
    if insertions:
        for after_id in sorted(insertions.keys(), reverse=True):
            for ins in reversed(insertions[after_id]):
                new_text = ins.get('revised_text', '')
                if new_text:
                    # Map float after_id to 1-based paragraph ID string
                    # rebuild uses 0-based indexing, inserter uses 1-based "p_N" IDs
                    para_id_str = f"p_{int(after_id) + 1}"
                    # Split on double newlines for multi-paragraph insertions
                    paragraphs = new_text.split('\n\n')
                    # Insert in reverse so multi-paragraph inserts stay in order
                    for p_text in reversed(paragraphs):
                        p_text = p_text.strip()
                        if p_text:
                            insert_paragraph_after(doc, para_id_str, p_text)

    doc.save(output_path)
    print(f"Saved revised document to: {output_path}")
    return True


def main():
    if len(sys.argv) < 4:
        print("Usage: rebuild_docx.py <original.docx> <output.docx> <revision1.json> [revision2.json ...]")
        sys.exit(1)

    original_path = sys.argv[1]
    output_path = sys.argv[2]
    json_paths = sys.argv[3:]

    print(f"Loading revisions from {len(json_paths)} file(s)...")
    revisions = load_revisions(*json_paths)
    print(f"Loaded {len(revisions)} revision entries")

    # Separate revisions by type
    modify_count = sum(1 for r in revisions.values() if r.get('action') in ('revise', 'replace'))
    insert_count = sum(1 for r in revisions.values() if r.get('action') == 'insert_after')
    delete_count = sum(1 for r in revisions.values() if r.get('action') == 'delete')
    print(f"  Revisions: {modify_count}, Insertions: {insert_count}, Deletions: {delete_count}")

    rebuild_document(original_path, output_path, revisions)


if __name__ == '__main__':
    main()
