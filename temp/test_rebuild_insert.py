"""Quick validation: rebuild_docx inserts paragraphs in correct position."""
import json
import sys
import tempfile
from pathlib import Path
from docx import Document

# Add paths
sys.path.insert(0, str(Path(__file__).parent.parent / "docx-tools"))

# Create test document
doc = Document()
doc.add_paragraph("Section 1. First provision.")
doc.add_paragraph("Section 2. Second provision.")
doc.add_paragraph("Section 3. Third provision.")
test_doc = Path(tempfile.mkdtemp()) / "test_input.docx"
doc.save(str(test_doc))

# Create revision with an insert_after
revisions = {
    "1": {
        "action": "insert_after",
        "insert_after_id": 0,
        "revised_text": "Section 1.1. New provision inserted after Section 1.",
    }
}
rev_file = test_doc.parent / "revisions.json"
rev_file.write_text(json.dumps(revisions))

# Import and run rebuild
sys.path.insert(0, str(Path(__file__).parent))
import rebuild_docx
output = test_doc.parent / "output.docx"
loaded = rebuild_docx.load_revisions(str(rev_file))
rebuild_docx.rebuild_document(str(test_doc), str(output), loaded)

# Verify
result = Document(str(output))
texts = [p.text for p in result.paragraphs if p.text.strip()]
print("Paragraphs after rebuild:")
for i, t in enumerate(texts):
    print(f"  {i}: {t}")

assert len(texts) == 4, f"Expected 4 paragraphs, got {len(texts)}"
assert "Section 1.1" in texts[1], f"Insert should be at position 1, got: {texts[1]}"
print("\nPASS: Insertion at correct position!")
