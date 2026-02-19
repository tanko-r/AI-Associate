"""
Write content to .docx files.

Converts markdown-formatted text or plain text into a properly formatted
Word document with headings, bold, italic, table, and list support.
"""

import re
from pathlib import Path
from typing import Optional, List

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def write_docx(
    content: str,
    output_path: str,
    template_path: Optional[str] = None,
) -> str:
    """
    Write content to a .docx file.

    Args:
        content: Markdown-formatted text or plain text to write.
        output_path: Path for the output .docx file.
        template_path: Optional path to a .docx template for style inheritance.

    Returns:
        The output file path.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if template_path and Path(template_path).exists():
        doc = Document(str(template_path))
        # Clear template content but keep styles
        for para in list(doc.paragraphs):
            p_element = para._element
            p_element.getparent().remove(p_element)
    else:
        doc = Document()
        _apply_default_styles(doc)

    _render_markdown_to_docx(doc, content)
    doc.save(str(output_path))
    return str(output_path)


def _apply_default_styles(doc: Document):
    """Apply clean professional defaults to a new document."""
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.bold = True
        if level == 1:
            heading_style.font.size = Pt(16)
            heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            heading_style.font.size = Pt(14)
        elif level == 3:
            heading_style.font.size = Pt(12)


def _render_table(doc: Document, headers: List[str], rows: List[List[str]], style: str = None):
    """Render a table from structured data.

    Args:
        doc: Document object to add table to.
        headers: List of header strings.
        rows: List of row lists (each row is a list of cell strings).
        style: Optional table style name.

    Returns:
        The table object.
    """
    table = doc.add_table(rows=1, cols=len(headers))

    # Apply table style
    if style:
        try:
            table.style = style
        except KeyError:
            try:
                table.style = 'Table Grid'
            except KeyError:
                pass
    else:
        try:
            table.style = 'Table Grid'
        except KeyError:
            pass

    # Header row -- bold text
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ''
        para = hdr_cells[i].paragraphs[0]
        _add_inline_markdown(para, header.strip())
        for run in para.runs:
            run.bold = True

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = ''
                para = row_cells[i].paragraphs[0]
                _add_inline_markdown(para, str(cell_text).strip())

    return table


def _parse_markdown_table(lines: List[str], start: int) -> tuple:
    """Parse a markdown table starting at the given line index.

    Args:
        lines: All lines of the content.
        start: Index of the first line of the table (header row).

    Returns:
        Tuple of (headers, rows, end_index) where end_index is the line
        after the last table line.
    """
    # Parse header row
    header_line = lines[start].strip()
    headers = [cell.strip() for cell in header_line.strip('|').split('|')]

    # Check for separator row
    end = start + 1
    if end < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[end].strip()):
        end += 1  # Skip separator row

    # Parse data rows
    rows = []
    while end < len(lines):
        line = lines[end].strip()
        if not line.startswith('|'):
            break
        # Skip if it looks like another separator
        if re.match(r'^\|[\s\-:|]+\|$', line):
            end += 1
            continue
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        rows.append(cells)
        end += 1

    return headers, rows, end


def _render_markdown_to_docx(doc: Document, content: str):
    """Parse markdown content and add it to the document."""
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Blank lines
        if not line.strip():
            i += 1
            continue

        # Markdown tables (lines starting with |)
        if line.strip().startswith('|'):
            headers, rows, end = _parse_markdown_table(lines, i)
            _render_table(doc, headers, rows)
            i = end
            continue

        # Ordered lists (1. item, 2. item, etc.)
        ol_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if ol_match:
            text = ol_match.group(2)
            try:
                para = doc.add_paragraph(style='List Number')
            except KeyError:
                # Fallback: add paragraph with number prefix preserved
                para = doc.add_paragraph()
                _add_inline_markdown(para, line.strip())
                i += 1
                continue
            _add_inline_markdown(para, text)
            i += 1
            continue

        # Unordered lists (- item or * item)
        ul_match = re.match(r'^[-*]\s+(.+)$', line.strip())
        if ul_match:
            text = ul_match.group(1)
            try:
                para = doc.add_paragraph(style='List Bullet')
            except KeyError:
                # Fallback: add paragraph with bullet prefix preserved
                para = doc.add_paragraph()
                _add_inline_markdown(para, line.strip())
                i += 1
                continue
            _add_inline_markdown(para, text)
            i += 1
            continue

        # Regular paragraph (may contain inline markdown)
        para = doc.add_paragraph()
        _add_inline_markdown(para, line.strip())
        i += 1


def _add_inline_markdown(paragraph, text: str):
    """Add text with inline markdown formatting (bold, italic) to a paragraph."""
    # Pattern to match **bold**, *italic*, or plain text segments
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    for match in re.finditer(pattern, text):
        full = match.group(0)
        if full.startswith('**') and full.endswith('**'):
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif full.startswith('*') and full.endswith('*'):
            run = paragraph.add_run(match.group(3))
            run.italic = True
        else:
            paragraph.add_run(match.group(4))
