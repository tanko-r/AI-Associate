"""Tests for contract analyzer."""
import pytest
from docx import Document


@pytest.fixture
def contract_docx(tmp_path):
    doc = Document()
    doc.add_heading("PURCHASE AND SALE AGREEMENT", level=1)
    doc.add_paragraph("This Purchase and Sale Agreement (this \"Agreement\") is entered into by Acme Corp (\"Seller\") and Widget Inc (\"Buyer\").")
    doc.add_paragraph("1.1 Seller represents and warrants that Seller has good and marketable title to the Property.")
    doc.add_paragraph("2.1 Buyer shall pay the Purchase Price of $1,000,000 at closing.")
    doc.add_paragraph("3.1 Seller shall indemnify and hold harmless Buyer from any claims arising from a breach of Seller's representations.")
    doc.add_paragraph("4.1 Either party may terminate this Agreement upon thirty (30) days written notice if the other party is in material default.")
    path = tmp_path / "psa.docx"
    doc.save(str(path))
    return path


def test_analyze_returns_structured_output(contract_docx):
    from core.analyzer import analyze_contract
    result = analyze_contract(str(contract_docx), representation="buyer")
    assert "risk_categories" in result
    assert "provisions_by_concept" in result
    assert "representation" in result


def test_analyze_identifies_risk_categories(contract_docx):
    from core.analyzer import analyze_contract
    result = analyze_contract(str(contract_docx), representation="buyer")
    categories = [c["category"] for c in result["risk_categories"]]
    assert len(categories) > 0


def test_analyze_includes_representation(contract_docx):
    from core.analyzer import analyze_contract
    result = analyze_contract(str(contract_docx), representation="seller")
    assert result["representation"] == "seller"
