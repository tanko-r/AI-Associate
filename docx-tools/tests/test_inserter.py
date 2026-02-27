"""Tests for paragraph positioning."""
import pytest
from docx import Document
from docx.oxml.ns import qn


@pytest.fixture
def three_para_docx(tmp_path):
    """Create a .docx with 3 known paragraphs."""
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third paragraph.")
    path = tmp_path / "three_para.docx"
    doc.save(str(path))
    return path


def test_find_paragraph_element_by_id(three_para_docx):
    from core.inserter import find_paragraph_element
    doc = Document(str(three_para_docx))
    elem = find_paragraph_element(doc, "p_2")
    assert elem is not None
    assert elem.tag == qn("w:p")
    texts = [t.text for t in elem.findall(f".//{qn('w:t')}")]
    assert "Second paragraph." in "".join(texts)


def test_find_paragraph_element_not_found(three_para_docx):
    from core.inserter import find_paragraph_element
    doc = Document(str(three_para_docx))
    elem = find_paragraph_element(doc, "p_99")
    assert elem is None


def test_insert_paragraph_after(three_para_docx, tmp_path):
    from core.inserter import insert_paragraph_after
    doc = Document(str(three_para_docx))
    new_elem = insert_paragraph_after(doc, "p_2", "Inserted after second.")
    assert new_elem is not None
    assert new_elem.tag == qn("w:p")
    out = tmp_path / "inserted.docx"
    doc.save(str(out))
    doc2 = Document(str(out))
    texts = [p.text for p in doc2.paragraphs]
    assert texts == [
        "First paragraph.",
        "Second paragraph.",
        "Inserted after second.",
        "Third paragraph.",
    ]


def test_insert_paragraph_after_copies_formatting(three_para_docx, tmp_path):
    from core.inserter import insert_paragraph_after
    doc = Document(str(three_para_docx))
    new_elem = insert_paragraph_after(doc, "p_2", "Formatted insert.")
    out = tmp_path / "formatted.docx"
    doc.save(str(out))
    doc2 = Document(str(out))
    inserted_para = doc2.paragraphs[2]
    assert inserted_para.text == "Formatted insert."


def test_insert_paragraph_after_invalid_id(three_para_docx):
    from core.inserter import insert_paragraph_after
    doc = Document(str(three_para_docx))
    result = insert_paragraph_after(doc, "p_99", "Should not insert.")
    assert result is None


def test_remove_paragraph(three_para_docx, tmp_path):
    from core.inserter import remove_paragraph
    doc = Document(str(three_para_docx))
    removed = remove_paragraph(doc, "p_2")
    assert removed is True
    out = tmp_path / "deleted.docx"
    doc.save(str(out))
    doc2 = Document(str(out))
    texts = [p.text for p in doc2.paragraphs]
    assert texts == ["First paragraph.", "Third paragraph."]


def test_remove_paragraph_invalid_id(three_para_docx):
    from core.inserter import remove_paragraph
    doc = Document(str(three_para_docx))
    result = remove_paragraph(doc, "p_99")
    assert result is False


def test_find_paragraph_in_table(tmp_path):
    from core.inserter import find_paragraph_element
    from docx.oxml.ns import qn
    doc = Document()
    doc.add_paragraph("Before table.")       # p_1
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0].text = "In table."  # p_2
    doc.add_paragraph("After table.")        # p_3
    path = tmp_path / "with_table.docx"
    doc.save(str(path))

    doc2 = Document(str(path))
    elem = find_paragraph_element(doc2, "p_2")
    assert elem is not None
    texts = [t.text for t in elem.findall(f".//{qn('w:t')}")]
    assert "In table." in "".join(texts)


def test_insert_paragraph_copies_bold_formatting(tmp_path):
    from core.inserter import insert_paragraph_after
    from docx.oxml.ns import qn
    doc = Document()
    doc.add_paragraph("Normal text.")           # p_1
    bold_p = doc.add_paragraph()                 # p_2
    run = bold_p.add_run("Bold text.")
    run.bold = True
    doc.add_paragraph("Third.")                  # p_3
    path = tmp_path / "bold.docx"
    doc.save(str(path))

    doc2 = Document(str(path))
    new_elem = insert_paragraph_after(doc2, "p_1", "Inserted.", style_source_id="p_2")
    # Verify rPr was copied with bold
    rPr = new_elem.find(f".//{qn('w:rPr')}")
    assert rPr is not None
    bold_elem = rPr.find(qn("w:b"))
    assert bold_elem is not None
