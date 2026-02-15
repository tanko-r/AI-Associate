"""Tests for docx comparer."""
import pytest
from docx import Document


@pytest.fixture
def two_versions(tmp_path):
    """Create two versions of a document."""
    doc1 = Document()
    doc1.add_paragraph("The Seller shall deliver the Property on the Closing Date.")
    doc1.add_paragraph("The Purchase Price shall be One Million Dollars ($1,000,000).")
    doc1.add_paragraph("This Agreement may be terminated by either party.")
    path1 = tmp_path / "v1.docx"
    doc1.save(str(path1))

    doc2 = Document()
    doc2.add_paragraph("The Seller shall deliver the Property on or before the Closing Date.")
    doc2.add_paragraph("The Purchase Price shall be One Million Dollars ($1,000,000).")
    doc2.add_paragraph("This Agreement may be terminated by Buyer only upon thirty (30) days written notice.")
    path2 = tmp_path / "v2.docx"
    doc2.save(str(path2))
    return path1, path2


def test_compare_returns_diff_report(two_versions):
    from core.comparer import compare_docx
    path1, path2 = two_versions
    result = compare_docx(str(path1), str(path2))
    assert isinstance(result, dict)
    assert "summary" in result
    assert "changes" in result


def test_compare_detects_modifications(two_versions):
    from core.comparer import compare_docx
    path1, path2 = two_versions
    result = compare_docx(str(path1), str(path2))
    changes = result["changes"]
    modified = [c for c in changes if c["type"] == "modified"]
    assert len(modified) == 2  # paragraphs 1 and 3 changed


def test_compare_detects_no_changes(tmp_path):
    from core.comparer import compare_docx
    doc = Document()
    doc.add_paragraph("Identical content.")
    path1 = tmp_path / "a.docx"
    path2 = tmp_path / "b.docx"
    doc.save(str(path1))
    doc.save(str(path2))
    result = compare_docx(str(path1), str(path2))
    assert result["summary"]["total_changes"] == 0


def test_compare_with_redline_output(two_versions, tmp_path):
    from core.comparer import compare_docx
    path1, path2 = two_versions
    redline_path = tmp_path / "comparison.docx"
    result = compare_docx(str(path1), str(path2), redline_output=str(redline_path))
    assert redline_path.exists()
    assert "redline_path" in result
