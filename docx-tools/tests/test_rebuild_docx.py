"""Tests for rebuild_docx formatting preservation."""
import sys
import os
import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Ensure temp/ (at project root) is importable
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))


def _build_multi_run_docx(path, runs_spec):
    """Build a .docx with a single paragraph containing runs per runs_spec.

    Args:
        path: Path to save the .docx file.
        runs_spec: list of (text, bold) tuples.

    Returns:
        The saved path.
    """
    doc = Document()
    para = doc.add_paragraph()
    # Clear the default empty run
    for r in list(para._p.findall(qn('w:r'))):
        para._p.remove(r)

    for text, bold in runs_spec:
        r = OxmlElement('w:r')
        if bold:
            rpr = OxmlElement('w:rPr')
            rpr.append(OxmlElement('w:b'))
            r.append(rpr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        para._p.append(r)

    doc.save(str(path))
    return path


def test_rebuild_preserves_per_run_formatting(tmp_path):
    """Bold 'Seller' run survives rebuild when only the end of text changes."""
    from temp.rebuild_docx import rebuild_document

    input_path = _build_multi_run_docx(
        tmp_path / "input.docx",
        [("The ", False), ('"Seller"', True), (" shall deliver.", False)]
    )
    output_path = tmp_path / "output.docx"

    # Revision changes only the end -- bold "Seller" should survive
    revisions = {
        0.0: {
            "action": "revise",
            "revised_text": 'The "Seller" shall convey.',
        }
    }

    rebuild_document(str(input_path), str(output_path), revisions)

    doc = Document(str(output_path))
    p = doc.paragraphs[0]._p
    all_runs = p.findall(qn('w:r'))

    # Find run containing "Seller" -- must have <w:b>
    seller_found = False
    for run in all_runs:
        texts = [t.text for t in run.findall(qn('w:t'))]
        run_text = ''.join(t or '' for t in texts)
        if 'Seller' in run_text:
            rpr = run.find(qn('w:rPr'))
            assert rpr is not None, "Bold Seller run must have <w:rPr>"
            assert rpr.find(qn('w:b')) is not None, "Seller run must be bold"
            seller_found = True
        elif run_text.strip() and 'Seller' not in run_text:
            # Non-Seller runs should NOT be bold
            rpr = run.find(qn('w:rPr'))
            if rpr is not None:
                assert rpr.find(qn('w:b')) is None, f"Run '{run_text}' should not be bold"

    assert seller_found, "Should find a run containing 'Seller' with bold formatting"


def test_rebuild_longer_text_uses_last_run_format(tmp_path):
    """Characters beyond original length inherit the last run's formatting."""
    from temp.rebuild_docx import rebuild_document

    # Last run is bold
    input_path = _build_multi_run_docx(
        tmp_path / "input.docx",
        [("Normal start. ", False), ("Bold end.", True)]
    )
    output_path = tmp_path / "output.docx"

    # Revised text is significantly longer
    revisions = {
        0.0: {
            "action": "revise",
            "revised_text": "Normal start. Bold end. Plus extra text beyond original.",
        }
    }

    rebuild_document(str(input_path), str(output_path), revisions)

    doc = Document(str(output_path))
    p = doc.paragraphs[0]._p
    all_runs = p.findall(qn('w:r'))

    # The last run (containing excess text) should inherit bold from the last original run
    last_run = all_runs[-1]
    rpr = last_run.find(qn('w:rPr'))
    assert rpr is not None, "Excess-text run should inherit last run's rPr"
    assert rpr.find(qn('w:b')) is not None, "Excess text should inherit bold from last original run"


def test_rebuild_shorter_text_no_trailing_runs(tmp_path):
    """Shorter revised text produces no empty trailing runs."""
    from temp.rebuild_docx import rebuild_document

    input_path = _build_multi_run_docx(
        tmp_path / "input.docx",
        [("First run. ", False), ("Second run. ", True), ("Third run.", False)]
    )
    output_path = tmp_path / "output.docx"

    # Revised text is shorter -- only covers first run's worth
    revisions = {
        0.0: {
            "action": "revise",
            "revised_text": "Short.",
        }
    }

    rebuild_document(str(input_path), str(output_path), revisions)

    doc = Document(str(output_path))
    p = doc.paragraphs[0]._p
    all_runs = p.findall(qn('w:r'))

    # All runs should have non-empty text -- no empty trailing runs
    for run in all_runs:
        texts = [t.text for t in run.findall(qn('w:t'))]
        run_text = ''.join(t or '' for t in texts)
        assert len(run_text) > 0, "No empty trailing runs should exist"

    # Total text should match the revised text
    total_text = ''.join(
        t.text or '' for run in all_runs for t in run.findall(qn('w:t'))
    )
    assert total_text == "Short."
