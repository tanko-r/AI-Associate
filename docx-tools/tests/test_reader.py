"""Tests for docx reader."""
import pytest
from pathlib import Path
from docx import Document


@pytest.fixture
def sample_docx(tmp_path):
    """Create a minimal test .docx file."""
    doc = Document()
    doc.add_heading("PURCHASE AND SALE AGREEMENT", level=1)
    doc.add_paragraph("This Purchase and Sale Agreement (this \"Agreement\") is entered into as of January 1, 2026, by and between Acme Corp (\"Seller\") and Widget Inc (\"Buyer\").")
    doc.add_heading("Article I: Definitions", level=2)
    doc.add_paragraph("1.1 \"Property\" means the real property located at 123 Main Street.")
    doc.add_paragraph("1.2 \"Purchase Price\" means the sum of One Million Dollars ($1,000,000).")
    doc.add_heading("Article II: Sale and Purchase", level=2)
    doc.add_paragraph("2.1 Agreement to Sell.  Seller shall sell the Property to Buyer, and Buyer shall purchase the Property from Seller, for the Purchase Price, subject to the terms and conditions of this Agreement.")
    doc.add_paragraph("2.2 Closing Date.  The closing shall occur on or before March 1, 2026 (the \"Closing Date\").")
    path = tmp_path / "test_contract.docx"
    doc.save(str(path))
    return path


def test_read_docx_returns_dict(sample_docx):
    from core.reader import read_docx
    result = read_docx(str(sample_docx))
    assert isinstance(result, dict)
    assert "content" in result
    assert "metadata" in result
    assert "defined_terms" in result
    assert "sections" in result


def test_read_docx_extracts_paragraphs(sample_docx):
    from core.reader import read_docx
    result = read_docx(str(sample_docx))
    paragraphs = [item for item in result["content"] if item["type"] == "paragraph"]
    assert len(paragraphs) >= 5
    # Each paragraph has an ID
    for p in paragraphs:
        assert p["id"].startswith("p_")
        assert "text" in p


def test_read_docx_extracts_defined_terms(sample_docx):
    from core.reader import read_docx
    result = read_docx(str(sample_docx))
    terms = result["defined_terms"]
    assert "Agreement" in terms
    assert "Property" in terms
    assert "Purchase Price" in terms


def test_read_docx_extracts_sections(sample_docx):
    from core.reader import read_docx
    result = read_docx(str(sample_docx))
    assert len(result["sections"]) > 0


def test_read_docx_extracts_metadata(sample_docx):
    from core.reader import read_docx
    result = read_docx(str(sample_docx))
    assert "core_properties" in result["metadata"]


def test_read_docx_nonexistent_file():
    from core.reader import read_docx
    with pytest.raises(FileNotFoundError):
        read_docx("/nonexistent/file.docx")
