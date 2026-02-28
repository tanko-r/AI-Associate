"""Tests for docx redliner."""
import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


@pytest.fixture
def original_docx(tmp_path):
    """Create a test .docx with known content."""
    doc = Document()
    doc.add_paragraph("The Seller shall deliver the Property on the Closing Date.")
    doc.add_paragraph("The Purchase Price shall be One Million Dollars ($1,000,000).")
    doc.add_paragraph("This Agreement may be terminated by either party upon thirty (30) days written notice.")
    path = tmp_path / "original.docx"
    doc.save(str(path))
    return path


def test_redline_docx_creates_output(original_docx, tmp_path):
    from core.redliner import redline_docx
    revisions = {
        "p_1": {
            "original": "The Seller shall deliver the Property on the Closing Date.",
            "revised": "The Seller shall deliver the Property on or before the Closing Date."
        }
    }
    output_path = tmp_path / "redlined.docx"
    result = redline_docx(str(original_docx), revisions, str(output_path))
    assert Path(result).exists()


def test_redline_docx_with_no_changes(original_docx, tmp_path):
    from core.redliner import redline_docx
    output_path = tmp_path / "redlined.docx"
    result = redline_docx(str(original_docx), {}, str(output_path))
    assert Path(result).exists()


def test_redline_docx_preserves_unchanged(original_docx, tmp_path):
    from core.redliner import redline_docx
    revisions = {
        "p_1": {
            "original": "The Seller shall deliver the Property on the Closing Date.",
            "revised": "The Seller shall deliver the Property on or before the Closing Date."
        }
    }
    output_path = tmp_path / "redlined.docx"
    redline_docx(str(original_docx), revisions, str(output_path))
    doc = Document(str(output_path))
    # Paragraph 2 (index 1) should be unchanged
    assert "Purchase Price" in doc.paragraphs[1].text


def test_redline_insert_after(original_docx, tmp_path):
    """Insert a new paragraph after p_1 with tracked changes."""
    from core.redliner import redline_docx
    revisions = {
        "p_1": {
            "action": "insert_after",
            "text": "Seller shall provide all closing documents five (5) days prior to Closing.",
        }
    }
    output_path = tmp_path / "inserted.docx"
    result = redline_docx(str(original_docx), revisions, str(output_path))
    assert Path(result).exists()
    doc = Document(str(output_path))
    assert len(doc.paragraphs) == 4
    # Text inside <w:ins> is not visible via python-docx's .text property,
    # so we extract all text from the XML body (including tracked changes).
    all_xml_text = "".join(doc.element.body.itertext())
    assert "closing documents" in all_xml_text


def test_redline_insert_after_has_tracking_markup(original_docx, tmp_path):
    """Verify insert_after wraps content in w:ins."""
    from core.redliner import redline_docx
    from docx.oxml.ns import qn
    revisions = {
        "p_1": {
            "action": "insert_after",
            "text": "New provision here.",
        }
    }
    output_path = tmp_path / "tracked_insert.docx"
    redline_docx(str(original_docx), revisions, str(output_path))
    doc = Document(str(output_path))
    body = doc.element.body
    ins_elements = body.findall(f".//{qn('w:ins')}")
    assert len(ins_elements) > 0


def test_redline_delete_paragraph(original_docx, tmp_path):
    """Delete p_2 with tracked changes."""
    from core.redliner import redline_docx
    from docx.oxml.ns import qn
    revisions = {
        "p_2": {"action": "delete"}
    }
    output_path = tmp_path / "deleted.docx"
    result = redline_docx(str(original_docx), revisions, str(output_path))
    assert Path(result).exists()
    doc = Document(str(output_path))
    body = doc.element.body
    del_elements = body.findall(f".//{qn('w:del')}")
    assert len(del_elements) > 0


def test_redline_mixed_actions(original_docx, tmp_path):
    """Replace p_1, insert after p_2, delete p_3 — all in one pass."""
    from core.redliner import redline_docx
    revisions = {
        "p_1": {
            "original": "The Seller shall deliver the Property on the Closing Date.",
            "revised": "The Seller shall deliver the Property on or before the Closing Date.",
        },
        "p_2": {
            "action": "insert_after",
            "text": "Buyer acknowledges receipt of all disclosures.",
        },
        "p_3": {"action": "delete"},
    }
    output_path = tmp_path / "mixed.docx"
    result = redline_docx(str(original_docx), revisions, str(output_path))
    assert Path(result).exists()
    doc = Document(str(output_path))
    # 3 original + 1 inserted = 4 paragraphs (delete keeps para with markup)
    assert len(doc.paragraphs) == 4


def test_redline_backward_compatible(original_docx, tmp_path):
    """Existing format without 'action' key still works as replace."""
    from core.redliner import redline_docx
    revisions = {
        "p_1": {
            "original": "The Seller shall deliver the Property on the Closing Date.",
            "revised": "The Seller shall deliver the Property on or before the Closing Date."
        }
    }
    output_path = tmp_path / "compat.docx"
    result = redline_docx(str(original_docx), revisions, str(output_path))
    assert Path(result).exists()


# ----------------------------------------------------------------
# Formatting-preservation tests (multi-run paragraphs)
# ----------------------------------------------------------------

@pytest.fixture
def multi_run_docx(tmp_path):
    """Create a .docx with multi-run paragraphs for formatting tests.

    Paragraph 1: 'The "Seller" shall deliver the Property.'
      - Run 1 (normal): 'The '
      - Run 2 (bold):   '"Seller"'
      - Run 3 (normal): ' shall deliver the Property.'
    Paragraph 2: simple single-run paragraph for comparison.
    """
    doc = Document()

    para = doc.add_paragraph()
    # Clear the default empty run
    for r in list(para._p.findall(qn('w:r'))):
        para._p.remove(r)

    # Run 1: normal
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.set(qn('xml:space'), 'preserve')
    t1.text = 'The '
    r1.append(t1)
    para._p.append(r1)

    # Run 2: bold
    r2 = OxmlElement('w:r')
    rpr2 = OxmlElement('w:rPr')
    rpr2.append(OxmlElement('w:b'))
    r2.append(rpr2)
    t2 = OxmlElement('w:t')
    t2.set(qn('xml:space'), 'preserve')
    t2.text = '"Seller"'
    r2.append(t2)
    para._p.append(r2)

    # Run 3: normal
    r3 = OxmlElement('w:r')
    t3 = OxmlElement('w:t')
    t3.set(qn('xml:space'), 'preserve')
    t3.text = ' shall deliver the Property.'
    r3.append(t3)
    para._p.append(r3)

    # Paragraph 2: simple single-run paragraph
    doc.add_paragraph("The Purchase Price shall be One Million Dollars.")

    path = tmp_path / "multi_run.docx"
    doc.save(str(path))
    return path


def test_redline_preserves_bold_in_equal_segments(multi_run_docx, tmp_path):
    """Bold formatting on 'Seller' survives a redline that only changes text at the end."""
    from core.redliner import redline_docx

    revisions = {
        "p_1": {
            "original": 'The "Seller" shall deliver the Property.',
            "revised": 'The "Seller" shall deliver the Property on the Closing Date.',
        }
    }
    output_path = tmp_path / "bold_preserved.docx"
    redline_docx(str(multi_run_docx), revisions, str(output_path))
    doc = Document(str(output_path))
    p = doc.paragraphs[0]._p

    # Collect all <w:r> elements (including those inside <w:ins>/<w:del>)
    all_runs = p.findall(f".//{qn('w:r')}")

    # Find the run(s) containing "Seller" -- must have bold
    bold_found = False
    for run in all_runs:
        texts = [t.text for t in run.findall(qn('w:t')) + run.findall(qn('w:delText'))]
        run_text = ''.join(t or '' for t in texts)
        if '"Seller"' in run_text or 'Seller' in run_text:
            rpr = run.find(qn('w:rPr'))
            assert rpr is not None, "Bold run should have <w:rPr>"
            b_elem = rpr.find(qn('w:b'))
            assert b_elem is not None, "Seller run should be bold"
            bold_found = True

    assert bold_found, "Should find a run containing 'Seller' with bold formatting"

    # Verify normal runs (containing 'The ') are NOT bold
    for run in all_runs:
        texts = [t.text for t in run.findall(qn('w:t')) + run.findall(qn('w:delText'))]
        run_text = ''.join(t or '' for t in texts)
        if run_text.strip() == 'The':
            rpr = run.find(qn('w:rPr'))
            if rpr is not None:
                assert rpr.find(qn('w:b')) is None, "'The ' run should not be bold"


def test_redline_insertion_inherits_adjacent_formatting(multi_run_docx, tmp_path):
    """Inserted text adjacent to bold 'Seller' inherits bold formatting."""
    from core.redliner import redline_docx

    revisions = {
        "p_1": {
            "original": 'The "Seller" shall deliver the Property.',
            "revised": 'The "Seller" (the "Owner") shall deliver the Property.',
        }
    }
    output_path = tmp_path / "insert_format.docx"
    redline_docx(str(multi_run_docx), revisions, str(output_path))
    doc = Document(str(output_path))
    p = doc.paragraphs[0]._p

    # Find <w:ins> elements -- there should be at least one
    ins_elements = p.findall(f".//{qn('w:ins')}")
    assert len(ins_elements) > 0, "Should have tracked insertions"

    # The inserted text should inherit formatting from the adjacent bold run
    for ins in ins_elements:
        runs = ins.findall(qn('w:r'))
        for run in runs:
            texts = [t.text for t in run.findall(qn('w:t'))]
            run_text = ''.join(t or '' for t in texts)
            if 'Owner' in run_text or '(' in run_text:
                rpr = run.find(qn('w:rPr'))
                assert rpr is not None, "Inserted run adjacent to bold should have rPr"
                assert rpr.find(qn('w:b')) is not None, (
                    "Insertion should inherit bold from adjacent 'Seller' run"
                )


def test_redline_no_rpr_paragraph(tmp_path):
    """A paragraph with runs that have no <w:rPr> works correctly -- no crash, no spurious formatting."""
    from core.redliner import redline_docx

    doc = Document()
    para = doc.add_paragraph()
    for r in list(para._p.findall(qn('w:r'))):
        para._p.remove(r)

    # Add a run with NO rPr
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.set(qn('xml:space'), 'preserve')
    t1.text = 'Simple text with no formatting.'
    r1.append(t1)
    para._p.append(r1)

    input_path = tmp_path / "no_rpr.docx"
    doc.save(str(input_path))

    revisions = {
        "p_1": {
            "original": "Simple text with no formatting.",
            "revised": "Simple text with updated formatting.",
        }
    }
    output_path = tmp_path / "no_rpr_redlined.docx"
    redline_docx(str(input_path), revisions, str(output_path))
    doc2 = Document(str(output_path))
    p = doc2.paragraphs[0]._p
    all_runs = p.findall(f".//{qn('w:r')}")

    # Runs from style-inherited paragraph should have no <w:rPr>
    for run in all_runs:
        rpr = run.find(qn('w:rPr'))
        assert rpr is None, "Style-inherited run should not gain spurious <w:rPr>"


def test_redline_deletion_preserves_run_formatting(multi_run_docx, tmp_path):
    """Deleted 'Seller' text retains its bold formatting in the deletion markup."""
    from core.redliner import redline_docx

    revisions = {
        "p_1": {
            "original": 'The "Seller" shall deliver the Property.',
            "revised": 'The shall deliver the Property.',
        }
    }
    output_path = tmp_path / "del_format.docx"
    redline_docx(str(multi_run_docx), revisions, str(output_path))
    doc = Document(str(output_path))
    p = doc.paragraphs[0]._p

    # Find <w:del> elements
    del_elements = p.findall(f".//{qn('w:del')}")
    assert len(del_elements) > 0, "Should have tracked deletions"

    # At least one deletion run should have bold formatting (the "Seller" run)
    bold_deleted = False
    for del_elem in del_elements:
        for run in del_elem.findall(qn('w:r')):
            texts = [t.text for t in run.findall(qn('w:delText'))]
            run_text = ''.join(t or '' for t in texts)
            if 'Seller' in run_text:
                rpr = run.find(qn('w:rPr'))
                if rpr is not None:
                    b_elem = rpr.find(qn('w:b'))
                    if b_elem is not None:
                        bold_deleted = True

    assert bold_deleted, "Deleted 'Seller' run should preserve bold formatting"
