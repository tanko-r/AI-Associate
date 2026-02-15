"""Tests for docx structure extractor."""
import pytest
from docx import Document


@pytest.fixture
def contract_docx(tmp_path):
    """Create a contract-like .docx for structure extraction."""
    doc = Document()
    doc.add_heading("COMMERCIAL LEASE AGREEMENT", level=1)
    doc.add_paragraph("This Lease Agreement (this \"Lease\") is entered into by and between Big Landlord LLC (\"Landlord\") and Small Tenant Inc (\"Tenant\").")
    doc.add_heading("Article I: Premises", level=2)
    doc.add_paragraph("1.1 Landlord hereby leases to Tenant the premises located at 456 Commerce Drive (the \"Premises\").")
    doc.add_heading("Article II: Term", level=2)
    doc.add_paragraph("2.1 The term of this Lease shall commence on April 1, 2026 (the \"Commencement Date\") and shall expire on March 31, 2031.")
    doc.add_paragraph("2.2 Tenant shall have the option to renew this Lease for one additional five-year term upon written notice to Landlord not less than 180 days prior to expiration.")
    doc.add_heading("Article III: Rent", level=2)
    doc.add_paragraph("3.1 Tenant shall pay base rent of $5,000 per month (the \"Base Rent\"), due on the first day of each calendar month.")
    doc.add_paragraph("3.2 Base Rent shall increase by three percent (3%) on each anniversary of the Commencement Date.")
    path = tmp_path / "lease.docx"
    doc.save(str(path))
    return path


def test_extract_returns_structure(contract_docx):
    from core.extractor import extract_structure
    result = extract_structure(str(contract_docx))
    assert "sections" in result
    assert "defined_terms" in result
    assert "provisions" in result
    assert "exhibits" in result


def test_extract_finds_sections(contract_docx):
    from core.extractor import extract_structure
    result = extract_structure(str(contract_docx))
    section_titles = [s["title"] for s in result["sections"]]
    assert any("Premises" in t for t in section_titles)
    assert any("Term" in t for t in section_titles)
    assert any("Rent" in t for t in section_titles)


def test_extract_finds_defined_terms(contract_docx):
    from core.extractor import extract_structure
    result = extract_structure(str(contract_docx))
    term_names = [t["term"] for t in result["defined_terms"]]
    assert "Lease" in term_names
    assert "Premises" in term_names
    assert "Base Rent" in term_names


def test_extract_classifies_provisions(contract_docx):
    from core.extractor import extract_structure
    result = extract_structure(str(contract_docx))
    provision_types = set(p["type"] for p in result["provisions"])
    assert len(provision_types) > 0
