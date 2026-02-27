"""Tests for placeholder and reference scanner."""
import pytest
from docx import Document


@pytest.fixture
def contract_with_blanks(tmp_path):
    """Create a .docx with various placeholder patterns."""
    doc = Document()
    doc.add_paragraph("1. Purchase Price. The Purchase Price shall be $____________.")
    doc.add_paragraph("2. Closing Date. The closing shall occur on [enter date].")
    doc.add_paragraph("3. Deposit. Buyer shall deposit TBD with Escrow Agent.")
    doc.add_paragraph("4. Property Address. The Property is located at XXXX Main Street.")
    doc.add_paragraph("5. Inspection. See Exhibit A for inspection requirements.")
    doc.add_paragraph("6. Title. See Schedule 1 for permitted exceptions.")
    doc.add_paragraph("7. Normal paragraph with no placeholders at all.")
    # Signature block — underscores here should be EXCLUDED
    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement.")
    doc.add_paragraph("By: ________________")
    doc.add_paragraph("Name: ________________")
    doc.add_paragraph("Title: ________________")
    doc.add_paragraph("Date: ________________")
    path = tmp_path / "contract_blanks.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def contract_with_exhibits(tmp_path):
    """Create a .docx with exhibit references and definitions."""
    doc = Document()
    doc.add_paragraph("See Exhibit A for the property description.")
    doc.add_paragraph("The exceptions are listed in Exhibit B.")
    doc.add_paragraph("Refer to Schedule 1 for closing adjustments.")
    doc.add_paragraph("EXHIBIT A")
    doc.add_paragraph("Legal Description of the Property...")
    path = tmp_path / "exhibits.docx"
    doc.save(str(path))
    return path


def test_scan_placeholders_finds_blanks(contract_with_blanks):
    from core.scanner import scan_placeholders
    findings = scan_placeholders(str(contract_with_blanks))
    pattern_types = [f["pattern_type"] for f in findings]
    assert "entry_field" in pattern_types
    assert "tbd" in pattern_types
    assert "xx_placeholder" in pattern_types


def test_scan_placeholders_skips_signature_lines(contract_with_blanks):
    from core.scanner import scan_placeholders
    findings = scan_placeholders(str(contract_with_blanks))
    # None of the findings should be underscores from signature block
    sig_findings = [f for f in findings if f["pattern_type"] == "underscores"]
    # The $____________ in paragraph 1 contains underscores but should be caught as dollar_blank
    # The signature block underscores (By:, Name:, Title:, Date:) should NOT appear
    for f in sig_findings:
        assert "By:" not in f.get("text_context", "")
        assert "Name:" not in f.get("text_context", "")
        assert "Title:" not in f.get("text_context", "")
        assert "Date:" not in f.get("text_context", "")


def test_scan_placeholders_count(contract_with_blanks):
    from core.scanner import scan_placeholders
    findings = scan_placeholders(str(contract_with_blanks))
    # Should find: dollar blank OR underscores in para 1, entry field, TBD, XX placeholder = at least 4
    # Should NOT find: signature line underscores
    assert len(findings) >= 4


def test_scan_references_finds_missing(contract_with_exhibits):
    from core.scanner import scan_references
    report = scan_references(str(contract_with_exhibits))
    defined_names = [d for d in report["defined"]]
    missing_names = [m for m in report["missing"]]
    assert any("Exhibit A" in d for d in defined_names)
    assert any("Exhibit B" in m for m in missing_names) or any("B" in m for m in missing_names)
    assert any("Schedule" in m for m in missing_names) or any("1" in m for m in missing_names)


def test_scan_document_combined(contract_with_blanks):
    from core.scanner import scan_document
    report = scan_document(str(contract_with_blanks))
    assert "placeholders" in report
    assert "references" in report
    assert "summary" in report
    assert report["summary"]["placeholder_count"] > 0
