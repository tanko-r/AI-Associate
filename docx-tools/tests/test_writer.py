"""Tests for docx writer."""
import pytest
from pathlib import Path
from docx import Document


@pytest.fixture
def output_dir(tmp_path):
    return tmp_path


def test_write_docx_from_markdown(output_dir):
    from core.writer import write_docx
    content = """# Memorandum

**To:** Partner
**From:** Sara
**Date:** February 14, 2026
**Re:** Lease Termination Rights

## Question Presented

Whether the tenant may terminate the lease early under Section 5.2.

## Short Answer

Yes. The lease grants the tenant an early termination right upon 90 days' written notice after the third lease year.
"""
    output_path = output_dir / "memo.docx"
    write_docx(content, str(output_path))
    assert output_path.exists()
    doc = Document(str(output_path))
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "Memorandum" in full_text
    assert "Question Presented" in full_text


def test_write_docx_with_template(output_dir, tmp_path):
    from core.writer import write_docx
    # Create a template with custom font
    template = Document()
    style = template.styles['Normal']
    style.font.name = 'Times New Roman'
    template_path = tmp_path / "template.docx"
    template.save(str(template_path))

    content = "This is a test paragraph."
    output_path = output_dir / "from_template.docx"
    write_docx(content, str(output_path), template_path=str(template_path))
    assert output_path.exists()


def test_write_docx_creates_parent_dirs(output_dir):
    from core.writer import write_docx
    output_path = output_dir / "subdir" / "deep" / "memo.docx"
    write_docx("Test content", str(output_path))
    assert output_path.exists()


def test_write_docx_returns_path(output_dir):
    from core.writer import write_docx
    output_path = output_dir / "test.docx"
    result = write_docx("Content", str(output_path))
    assert result == str(output_path)
